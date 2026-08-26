"""内容提取器测试（M2.5 验收，设计文档第 4 节）。

- md/txt 直接读取（含非 utf-8 编码降级）；
- PDF：文字型提取成功；扫描件（无文字层）→ UnsupportedFormatError（unindexed）；
  损坏文件 → ExtractionFailedError（failed，可重试）；
- docx：段落 + 表格提取（python-docx 生成真实文件做回环）；
- zip/图片等不支持格式 → UnsupportedFormatError，不阻塞其他文件。
"""

import io

import pytest
from docx import Document

from app.domains.memory.extractors import (
    ExtractionFailedError,
    UnsupportedFormatError,
    extract_text,
)

def _build_pdf(objects: list[bytes]) -> bytes:
    """组装带正确 xref 表的最小 PDF（pypdf 要求 startxref 指向合法 xref）。"""
    out = bytearray(b"%PDF-1.4\n")
    offsets = []
    for i, obj in enumerate(objects, start=1):
        offsets.append(len(out))
        out += f"{i} 0 obj\n".encode() + obj + b"\nendobj\n"
    xref_pos = len(out)
    size = len(objects) + 1
    out += f"xref\n0 {size}\n".encode()
    out += b"0000000000 65535 f \n"
    for offset in offsets:
        out += f"{offset:010d} 00000 n \n".encode()
    out += f"trailer << /Size {size} /Root 1 0 R >>\nstartxref\n{xref_pos}\n%%EOF\n".encode()
    return bytes(out)


def _content_stream(stream: bytes) -> bytes:
    return b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n" + stream + b"\nendstream"


TEXT_PDF = _build_pdf(
    [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        _content_stream(b"BT /F1 24 Tf 100 700 Td (Hello Memory) Tj ET"),
    ]
)

# 无文字内容的"扫描件"PDF（空内容流）
SCAN_PDF = _build_pdf(
    [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R >>",
        _content_stream(b""),
    ]
)


def _make_docx() -> bytes:
    document = Document()
    document.add_paragraph("需求规格第一段。")
    document.add_paragraph("需求规格第二段。")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "表格单元甲"
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def test_markdown_and_txt() -> None:
    assert extract_text("guide.md", "# 标题\n\n正文。".encode()) == "# 标题\n\n正文。"
    assert extract_text("notes.txt", "纯文本内容。".encode()) == "纯文本内容。"


def test_txt_non_utf8_degrades_gracefully() -> None:
    """GBK 等非法 utf-8 字节不阻塞索引，降级 replace。"""
    text = extract_text("legacy.txt", "中文".encode("gbk"))
    assert len(text) > 0


def test_pdf_with_text_layer() -> None:
    text = extract_text("manual.pdf", TEXT_PDF)
    assert "Hello Memory" in text


def test_scanned_pdf_is_unsupported() -> None:
    with pytest.raises(UnsupportedFormatError):
        extract_text("scan.pdf", SCAN_PDF)


def test_corrupted_pdf_is_retryable_failure() -> None:
    with pytest.raises(ExtractionFailedError):
        extract_text("broken.pdf", b"not a pdf at all")


def test_docx_paragraphs_and_tables() -> None:
    text = extract_text("spec.docx", _make_docx())
    assert "需求规格第一段。" in text
    assert "需求规格第二段。" in text
    assert "表格单元甲" in text


def test_corrupted_docx_is_retryable_failure() -> None:
    with pytest.raises(ExtractionFailedError):
        extract_text("broken.docx", b"not a zip/docx")


def test_unsupported_formats() -> None:
    for filename in ("archive.zip", "photo.png", "avatar.jpg", "noext"):
        with pytest.raises(UnsupportedFormatError):
            extract_text(filename, b"data")
