"""将文档字节提取为供索引管道切块的纯文本。

`.md` 和 `.txt` 按 `UTF-8` 解码，失败时替换无效字符；`.pdf` 通过 `pypdf`
提取文字层，无文字层的扫描件视为不支持；`.docx` 通过 `python-docx` 提取段落和表格。

`UnsupportedFormatError` 表示格式不支持，对应不可重试的 `unindexed`；
`ExtractionFailedError` 表示支持的格式解析失败，对应可重试的 `failed`。
"""

import io

from pypdf import PdfReader
from pypdf.errors import PdfReadError

SUPPORTED_EXTENSIONS = frozenset({".md", ".txt", ".pdf", ".docx"})


class UnsupportedFormatError(Exception):
    """格式不支持读取内容，对应索引状态 `unindexed`。"""


class ExtractionFailedError(Exception):
    """支持的格式提取失败，对应可重试的索引状态 `failed`。"""


def _extract_text_file(data: bytes) -> str:
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError:
        # 替换无效字节，避免编码问题阻塞整份文档索引
        return data.decode("utf-8", errors="replace")


def _extract_pdf(data: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(data))
        pages = [(page.extract_text() or "") for page in reader.pages]
    except (PdfReadError, ValueError, OSError) as exc:
        raise ExtractionFailedError(f"PDF 解析失败: {exc}") from exc
    text = "\n\n".join(p for p in pages if p.strip())
    if not text.strip():
        # 无文字层的扫描件需要 `OCR`，当前提取器不支持
        raise UnsupportedFormatError("扫描件 PDF（无文字层）")
    return text


def _extract_docx(data: bytes) -> str:
    from docx import Document

    try:
        document = Document(io.BytesIO(data))
    except Exception as exc:  # `python-docx` 对损坏文件可能抛出多种异常
        raise ExtractionFailedError(f"DOCX 解析失败: {exc}") from exc
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells if cell.text.strip())
    return "\n\n".join(parts)


def extract_text(filename: str, data: bytes) -> str:
    """按文件扩展名提取纯文本。

    不支持的格式抛出 `UnsupportedFormatError`；解析失败抛出 `ExtractionFailedError`；
    空文档可以返回空字符串，由索引管道按空内容处理。
    """
    suffix = ("." + filename.rsplit(".", 1)[-1].lower()) if "." in filename else ""
    if suffix not in SUPPORTED_EXTENSIONS:
        raise UnsupportedFormatError(f"不支持读取内容的格式: {suffix or '<无扩展名>'}")
    if suffix in (".md", ".txt"):
        return _extract_text_file(data)
    if suffix == ".pdf":
        return _extract_pdf(data)
    return _extract_docx(data)
